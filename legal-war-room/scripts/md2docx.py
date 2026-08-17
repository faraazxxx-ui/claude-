#!/usr/bin/env python3
"""Markdown -> .docx via python-docx. Handles: h1-h4, tables, blockquotes, ul/ol, hr, bold/italic/code.
Usage: md2docx.py file.md [file2.md ...]"""
import re, sys, os
import docx
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

NAVY = RGBColor(0x21, 0x31, 0x5E)
MAROON = RGBColor(0x8A, 0x1F, 0x1F)
GREY = RGBColor(0x40, 0x43, 0x4B)

INLINE = re.compile(r"(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`|\[.+?\]\(.+?\))")

def add_runs(par, text):
    for tok in INLINE.split(text):
        if not tok: continue
        if tok.startswith("**") and tok.endswith("**"):
            r = par.add_run(tok[2:-2]); r.bold = True
        elif tok.startswith("*") and tok.endswith("*") and len(tok) > 2:
            r = par.add_run(tok[1:-1]); r.italic = True
        elif tok.startswith("`") and tok.endswith("`"):
            r = par.add_run(tok[1:-1]); r.font.name = "Consolas"; r.font.size = Pt(9)
        elif tok.startswith("[") and "](" in tok:
            label = tok[1:tok.index("](")]
            r = par.add_run(label); r.underline = True
        else:
            par.add_run(tok)

def table_rows(lines, i):
    rows = []
    while i < len(lines) and lines[i].strip().startswith("|"):
        cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
        if not all(re.fullmatch(r":?-{2,}:?", c) for c in cells):
            rows.append(cells)
        i += 1
    return rows, i

def convert(path):
    doc = docx.Document()
    st = doc.styles["Normal"]; st.font.name = "Calibri"; st.font.size = Pt(10)
    for sec in doc.sections:
        sec.page_width, sec.page_height = Inches(8.5), Inches(11)
        sec.left_margin = sec.right_margin = Inches(0.7)
        sec.top_margin = sec.bottom_margin = Inches(0.7)
    lines = open(path).read().split("\n")
    i = 0
    while i < len(lines):
        ln = lines[i]
        s = ln.strip()
        if s.startswith("|"):
            rows, i = table_rows(lines, i)
            if rows:
                ncol = max(len(r) for r in rows)
                t = doc.add_table(rows=len(rows), cols=ncol)
                t.style = "Light Grid Accent 1"
                for ri, row in enumerate(rows):
                    for ci in range(ncol):
                        cell = t.cell(ri, ci)
                        cell.paragraphs[0].text = ""
                        add_runs(cell.paragraphs[0], row[ci] if ci < len(row) else "")
                        for p in cell.paragraphs:
                            for r in p.runs:
                                r.font.size = Pt(8.5)
                                if ri == 0: r.bold = True
                doc.add_paragraph()
            continue
        if s.startswith("#"):
            level = len(s) - len(s.lstrip("#"))
            h = doc.add_heading("", level=min(level, 4))
            add_runs(h, s.lstrip("# "))
            for r in h.runs: r.font.color.rgb = MAROON if level == 1 else NAVY
        elif s.startswith(">"):
            p = doc.add_paragraph(style="Intense Quote")
            add_runs(p, s.lstrip("> "))
        elif re.match(r"^(-{3,}|\*{3,})$", s):
            p = doc.add_paragraph(); p.add_run("—" * 40).font.color.rgb = RGBColor(0xC9, 0xD2, 0xE4)
        elif re.match(r"^[-*] ", s):
            p = doc.add_paragraph(style="List Bullet"); add_runs(p, s[2:])
        elif re.match(r"^\d+\. ", s):
            p = doc.add_paragraph(style="List Number"); add_runs(p, re.sub(r"^\d+\. ", "", s))
        elif s:
            p = doc.add_paragraph(); add_runs(p, s)
        i += 1
    out = os.path.splitext(path)[0] + ".docx"
    doc.save(out)
    print(f"OK {out} ({os.path.getsize(out)//1024}KB)")

if __name__ == "__main__":
    for f in sys.argv[1:]:
        convert(f)
